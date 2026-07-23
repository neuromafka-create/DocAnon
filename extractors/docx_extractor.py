from __future__ import annotations

from pathlib import Path

from docx import Document

from core.models import ExtractionResult
from extractors.base import TextExtractor


class DocxExtractor(TextExtractor):
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def extract(self, file_path: Path) -> ExtractionResult:
        doc = Document(str(file_path))
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        return ExtractionResult(
            text=text,
            source_format="docx",
            metadata={"file_name": file_path.name, "paragraphs": len(doc.paragraphs)},
        )
