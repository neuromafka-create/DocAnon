from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from core.batch_pipeline import BatchPipeline
from core.config import AnonymizerConfig
from core.mapping import MappingStore
from core.models import AnonymizedDocument, Entity
from core.pipeline import AnonymizationPipeline, apply_placeholders
from exporters.docx_exporter import DocxExporter


class TestD1SharedPipeline:
    def test_batch_reuses_pipeline_instance(self) -> None:
        cfg = AnonymizerConfig()
        single = AnonymizationPipeline(cfg)
        batch = BatchPipeline(pipeline=single)
        assert batch._pipeline is single
        assert batch.config is single.config

    def test_apply_placeholders(self) -> None:
        text = "AAA and BBB"
        entities = [
            Entity("AAA", "X", 0, 3, 0.9),
            Entity("BBB", "Y", 8, 11, 0.9),
        ]
        out = apply_placeholders(
            text, entities, lambda o, lab: f"<{lab}>"
        )
        assert out == "<X> and <Y>"

    def test_register_and_anonymize_with_store(self, tmp_path: Path) -> None:
        pipe = AnonymizationPipeline(AnonymizerConfig())
        store = MappingStore()
        text = "Контакт: Иванов, ИНН 7707083893"
        entities = pipe.analyze(text)
        pipe.register_entities(text, entities, store, source_file="a.txt")
        anon = pipe.anonymize_with_store(text, entities, store)
        assert "7707083893" not in anon
        assert "<" in anon


class TestD4DocxExporter:
    def test_export_from_text(self, tmp_path: Path) -> None:
        doc = AnonymizedDocument(
            original_text="Hello\nWorld",
            anonymized_text="Hello\n<PERSON>",
            source_file="x.txt",
        )
        out = DocxExporter().export(doc, tmp_path / "out")
        assert out.exists()
        loaded = Document(str(out))
        texts = [p.text for p in loaded.paragraphs]
        assert any("<PERSON>" in t for t in texts)

    def test_roundtrip_preserves_paragraphs(self, tmp_path: Path) -> None:
        src = tmp_path / "src.docx"
        document = Document()
        document.add_paragraph("Сторона: ООО «Тест»")
        document.add_paragraph("Контакт: Иванов Иван")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "ФИО"
        table.rows[0].cells[1].text = "Петров"
        document.save(str(src))

        mapping = {
            "ООО «Тест»": "<ORG>",
            "Иванов Иван": "<PERSON>",
            "Петров": "<PERSON_2>",
        }
        out = tmp_path / "anon.docx"
        DocxExporter().export_from_source(src, mapping, out)

        loaded = Document(str(out))
        all_text = "\n".join(p.text for p in loaded.paragraphs)
        for row in loaded.tables[0].rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

        assert "ООО «Тест»" not in all_text
        assert "<ORG>" in all_text
        assert "Иванов Иван" not in all_text
        assert "<PERSON>" in all_text
        assert "Петров" not in all_text
        assert "<PERSON_2>" in all_text

    def test_export_uses_source_path(self, tmp_path: Path) -> None:
        src = tmp_path / "contract.docx"
        document = Document()
        document.add_paragraph("Email: a@b.ru")
        document.save(str(src))

        adoc = AnonymizedDocument(
            original_text="Email: a@b.ru",
            anonymized_text="Email: <EMAIL>",
            mapping={"a@b.ru": "<EMAIL>"},
            source_file=src.name,
            source_path=str(src),
        )
        out = DocxExporter().export(adoc, tmp_path / "result")
        loaded = Document(str(out))
        assert any("<EMAIL>" in p.text for p in loaded.paragraphs)


class TestSettingsBuildConfig:
    def test_settings_panel_importable(self) -> None:
        # headless-friendly: only import/module structure
        from gui.settings_panel import SettingsPanel, _SETTINGS_ENTITY_ORDER

        assert "PERSON" in _SETTINGS_ENTITY_ORDER
        assert "RU_INN" in _SETTINGS_ENTITY_ORDER
