from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from core.models import AnonymizedDocument
from core.xlsx_cells import apply_mapping_to_text
from exporters.base import Exporter

logger = logging.getLogger(__name__)


def _set_paragraph_text_keep_style(paragraph: Paragraph, new_text: str) -> None:
    """Заменить текст параграфа, сохранив стиль первого run (грубо)."""
    if paragraph.text == new_text:
        return

    # сохранить formatting первого run
    style_run = None
    if paragraph.runs:
        style_run = paragraph.runs[0]

    # очистить runs
    for run in paragraph.runs:
        run.text = ""

    if not paragraph.runs:
        run = paragraph.add_run(new_text)
        return

    # первый run — весь новый текст
    paragraph.runs[0].text = new_text
    if style_run is not None and paragraph.runs[0] is not style_run:
        # already first
        pass


def _replace_in_paragraph(paragraph: Paragraph, mapping: dict[str, str]) -> bool:
    full = paragraph.text
    if not full or not mapping:
        return False
    new = apply_mapping_to_text(full, mapping)
    if new == full:
        return False
    _set_paragraph_text_keep_style(paragraph, new)
    return True


def _iter_all_paragraphs(document: Document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    # headers / footers
    for section in document.sections:
        for p in section.header.paragraphs:
            yield p
        for p in section.footer.paragraphs:
            yield p
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


class DocxExporter(Exporter):
    """D4: экспорт DOCX.

    - если есть исходный ``.docx`` (source_path) — round-trip с заменой
      в параграфах/таблицах (стиль абзаца и первый run сохраняются);
    - иначе — новый документ из обезличенного текста.
    """

    def export(self, doc: AnonymizedDocument, output_path: Path) -> Path:
        output_path = Path(output_path).with_suffix(".docx")
        src = Path(doc.source_path) if doc.source_path else None

        if src and src.exists() and src.suffix.lower() == ".docx":
            return self.export_from_source(src, doc.mapping, output_path)

        return self.export_from_text(doc.anonymized_text, output_path)

    def export_from_source(
        self,
        source_path: Path,
        mapping: dict[str, str],
        output_path: Path,
    ) -> Path:
        output_path = Path(output_path).with_suffix(".docx")
        document = Document(str(source_path))
        changed = 0
        for paragraph in _iter_all_paragraphs(document):
            if _replace_in_paragraph(paragraph, mapping):
                changed += 1
        logger.info("DOCX round-trip: %d paragraphs updated", changed)
        document.save(str(output_path))
        return output_path

    def export_from_text(self, text: str, output_path: Path) -> Path:
        output_path = Path(output_path).with_suffix(".docx")
        document = Document()

        # разбить на абзацы; пустые строки — пустой paragraph
        lines = text.split("\n")
        for line in lines:
            p = document.add_paragraph(line if line.strip() else "")
            # простой эвристический «заголовок»
            if line.isupper() and 3 < len(line) < 80:
                try:
                    p.style = "Heading 1"
                except Exception:
                    pass
            elif re.match(r"^(#{1,3}\s|ДОГОВОР|Договор|Отчет|ОТЧЕТ)", line):
                try:
                    p.style = "Heading 2"
                except Exception:
                    pass

        document.save(str(output_path))
        return output_path
